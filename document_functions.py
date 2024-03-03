# Import modules
from docx import Document
import os
from openai_functions import chat_completion

# Length of document_templates & document_templates_names_for_saving should be same

# File names for saving respective template
document_templates_names_for_saving = ["lesson_plan", 
                                       "assessment", 
                                       "answers"]

# Load template of all documents
def document_templates():
    return [Document(f"{os.path.abspath('sample_templates/LESSON_PLAN_TEMPLATE.docx')}"), 
                      Document(f"{os.path.abspath('sample_templates/ASSESSMENT_TEMPLATE.docx')}"), 
                      Document(f"{os.path.abspath('sample_templates/MARKING_GUIDE_TEMPLATE.docx')}")]

def get_all_data_from_file(file_name):
    """
    This function loads a document from a file.

    Parameters:
        file_name (str): The name of the file to load.

    Returns:
        Document: The loaded document.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    try:
        return Document(f"{file_name}.docx")
    except FileNotFoundError:
        raise

def get_table_id_of_days(lesson_plan_days, doc_template):
    """
    This function takes a list of days and a document template and returns a list of table objects that contain those days.

    Parameters:
        lesson_plan_days (list): A list of days that are present in the lesson plan template.
        doc_template (Document): A document template that contains the lesson plan.

    Returns:
        list: A list of table objects that contain the days in the lesson plan.
        Error_error:  If there is an error while finding the table objects.
    """
    try:
        list_of_days_table_id = []
        for lesson_plan_day in lesson_plan_days:
            for i in range(0, len(doc_template.tables)):
                if  doc_template.tables[i].rows[0].cells[0].text == lesson_plan_day:
                    list_of_days_table_id.append(doc_template.tables[i])
                    break
        return list_of_days_table_id
    except Exception as e:
        return f"Error_{e}"

def get_table_id_from_template_contain_days():
    """ 
    This function give table_objects of those tables which contain days names
    
    Returns:
        list of those table_objects which contain days names
        Error_error -  if there is any problem to get table id from template contain days
    """
    try:
        days_names = ["MONDAY", "TUESDAY", "WEDNESDAY", 
                      "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY"]
        document_templates_local = document_templates()
        tables_id_of_tables_contain_days_name = []
        for i in range(0, len(document_templates_local[0].tables)):
            for day in days_names:
                if  str.lower(document_templates_local[0].tables[i].rows[0].cells[0].text) == str.lower(day):
                    tables_id_of_tables_contain_days_name.append(document_templates_local[0].tables[i])
                    break
        return tables_id_of_tables_contain_days_name
    except Exception as e:
        return f"Error_{e}"

def get_available_days_name():
    """ 
    This function get days name which are available in provided template
    
    Returns:
        list of days available in template
        Error_error -  if there is any problem to get names of days from template
    """
    try:
        days_available_in_template = []
        for table_id in get_table_id_from_template_contain_days():
            days_available_in_template.append(table_id.rows[0].cells[0].text)
        return days_available_in_template
    except Exception as e:
        return f"Error_{e}"

def update_intro_table(teacher_name, course, unit_title, week):
    """
    This function will update template of first table which is introduction table
    
    Parameters (required):
        - teacher_name, course, unit and title, week
    
    Returns:
        True - if table updated successfully
        Error_error -  if table is not updated for any reason
    """
    try:
        # First table about Introduction
        # for index_num, template in enumerate(document_templates):
        #     intro_table = template.tables[0]
        #     intro_table.rows[0].cells[1].text = teacher_name
        #     intro_table.rows[0].cells[3].text = course
        #     intro_table.rows[1].cells[1].text = unit_title
        #     intro_table.rows[1].cells[3].text = week
        #     save_document(template, document_templates_names_for_saving[index_num])
        global document_templates_global
        document_templates_global = document_templates()

        intro_table = document_templates_global[0].tables[0]
        intro_table.rows[0].cells[1].text = teacher_name
        intro_table.rows[0].cells[3].text = course
        intro_table.rows[1].cells[1].text = unit_title
        intro_table.rows[1].cells[3].text = week
        save_document(document_templates_global[0], document_templates_names_for_saving[0])
        return True, document_templates_global[0]
    except Exception as e:
        return f"Error_{e}"

# =================================================
# Function to add key as bold and value as normal
def add_key_value(key, value, table_id):
    p = table_id.add_paragraph()
    p.add_run(key).bold = True
    p.add_run(": " + value)

# Function to handle list of dictionaries
def handle_list_of_dicts(key, value, table_id):
    index = 1
    add_key_value(key, "", table_id)
    for item in value:
        for inner_key, inner_value in item.items():
            p = table_id.add_paragraph()
            p.add_run(f"Stage {index} - {inner_key}").bold = True
            p.add_run(": " + inner_value)
            index += 1

def adjust_lesson_body(table_id, lesson_body, time_duration_mints, activity):
    prompt = f"""
    Your task is to generate a lesson plan with activity of provided content, delimited by triple 
    backticks.
    Exclude 26% of time. Remaining time should be divided into stages.
    On each stage we divide the time accodingly so we can show the time to be spent on each stage.
    If the time is 90 or more, let us make it around 6 or more stages but not more than 8.\
    If the time is around 60 make it around 5 stages, \
    If the time is 40 we make it 3-4 stages and if the time is below 40 we keep then at 3.
    Given activity should connect to the lesson plan and you should create the materials needed and the activity should embedded in the lesson plan.
    Be carefull with number of stages, time specified on each stage and connection of activity with lesson plan.
    
    Do not include any explanations, only provide RFC8259 compliant JSON response following this format without deviation.
    {{"Lesson_Title": "title of lesson",
    "Duration": "{time_duration_mints} minutes",
    "Focus": "Describe main focus during lecture. It should be 2 to 3 lines",
    "Materials":"should be separted and started by new line",
    "Activity": "{activity}",
    "Lesson_Stages": [{{"Stage name (minutes spend on each stage)":"Stage details"}}],
    "}}
    content: ```{lesson_body}```
    """

    lesson_body = chat_completion(prompt)

    for key, value in lesson_body.items():
        if isinstance(value, list):  # If value is a list of dictionaries
            handle_list_of_dicts(key, value, table_id)
        else:
            add_key_value(key, value, table_id)
# =================================================
def update_table_for_lesson_plan(list_of_table_id, list_of_gpt_response, list_of_activity):
    """
    Update the table of lesson plan with the given list of table id and gpt response.

    Parameters: 
        list_of_table_id (list): list of table id
        list_of_gpt_response (list): list of gpt response
    
    Returns:
        True if document is saved successfully, else returns an error message.
    """
    try:
        # Table of key concept and update it's value
        terminology = ""
        for index_number, table_id in enumerate(list_of_table_id):
            # list_of_gpt_response is a list of dictionaries
            # Check if index_number is within the range of list_of_gpt_response
            if index_number < len(list_of_gpt_response) :
                terminology += list_of_gpt_response[index_number]["terminology"]+"\n"

                # table_id has rows and cells structure
                # table_id.rows[1].cells[1].text = "\n".join(["SWBAT "+line for line in list_of_gpt_response[index_number]["aims_and_objective"].split("\n")])     # Must add SWBAT in this field
                table_id.rows[1].cells[1].text = list_of_gpt_response[index_number]["aims_and_objective"]
                table_id.rows[2].cells[1].text = list_of_gpt_response[index_number]["introduction"]
                # table_id.rows[3].cells[1].text = list_of_gpt_response[index_number]["lesson_body"]
                specific_table_id = table_id.rows[3].cells[1]
                adjust_lesson_body(specific_table_id, list_of_gpt_response[index_number]["lesson_body"], 
                                list_of_gpt_response[index_number]["Duration"], 
                                list_of_activity[index_number])
                table_id.rows[4].cells[1].text = list_of_gpt_response[index_number]["conclusion"]

        # Update key concept and terminology table
        document_templates_global[0].tables[1].rows[1].cells[0].text = terminology

        # Save the document
        return save_document(document_templates_global[0], document_templates_names_for_saving[0])
    except Exception as e:
        return f"Error_{e}"

# =================== Not in use for now SECTION START ===================
def update_table_for_assessment_and_marking_guide(list_of_table_id, list_of_gpt_response, file_description):
    """
    Update the table of assessment and marking guide with the given list of table id and gpt response.

    Parameters:
        list_of_table_id (list): list of table id
        list_of_gpt_response (list): list of gpt response
        file_description (str): description of file (assessment or marking guide)

    Returns:
        True if document is saved successfully, else returns an error message.
    """
    try:
        # Table of key concept and update it's value
        for index_number, table_id in enumerate(list_of_table_id):
            # list_of_gpt_response is a list of dictionaries
            # Check if index_number is within the range of list_of_gpt_response
            if index_number < len(list_of_gpt_response):
                # table_id has rows and cells structure
                if file_description == "assessment":
                    table_id.rows[1].cells[1].text = list_of_gpt_response[index_number]["assessment"]

                elif file_description == "marking_guide":
                    table_id.rows[1].cells[1].text = list_of_gpt_response[index_number]["marking_guide"]

        # Save the document
        if file_description == "assessment":
            return save_document(document_templates()[1], document_templates_names_for_saving[1])
        elif file_description == "marking_guide":
            return save_document(document_templates()[2], document_templates_names_for_saving[2])
        
    except Exception as e:
        return f"Error_{e}"
# =================== Not in use for now SECTION END =====================

def update_assessment_and_marking_guide(document_template, list_of_gpt_response, file_description):
    try:
        formatted_response = ""
        for dictionary in list_of_gpt_response:
            if dictionary[file_description].endswith("\n"):
                formatted_response += dictionary[file_description]
            else:
                formatted_response += dictionary[file_description] + "\n"
        # formatted_response = formatted_response.replace("\n", "\n\n")
        formatted_response = formatted_response.strip().split("\n")
        for i, data in enumerate(formatted_response, start=1):
            document_template.add_paragraph(f"{i}. {data}")

        # Save the document
        if file_description == "assessment":
            return save_document(document_template, document_templates_names_for_saving[1])
        elif file_description == "answers":
            return save_document(document_template, document_templates_names_for_saving[2])
    except Exception as e:
        return f"Error_{e}"

def save_document(template, name_of_document):
    """
    This function will save document

    Parameters:
        template (Document): The document to save.
        name_of_document (str): The name of the file to save the document as.

    Returns:
        bool: True if the document was saved successfully
        Error_error -  if table is not updated for any reason
    """
    try:
        template.save(f"{name_of_document}.docx")
        return True
    except Exception as e:
        return f"Error_{e}"
